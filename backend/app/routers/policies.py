"""
Company internal policy documents — upload, ingest into the RAG index, and
ask grounded questions against them (discount rules, approval thresholds,
technical standards, compliance clauses, etc.). Tenant-wide, not project-scoped.
"""
import io
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from pypdf import PdfReader
from docx import Document as DocxDocument

from .. import models, schemas
from ..database import get_db
from ..auth import get_current_user
from ..services import s3_service
from ..services.policy_rag import ingest_policy_document, answer_policy_question

router = APIRouter(prefix="/policies", tags=["policies"])


def _extract_text(storage_path: str, filename: str) -> str:
    raw = s3_service.read_file(storage_path)
    if filename.lower().endswith(".pdf"):
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if filename.lower().endswith(".docx"):
        document = DocxDocument(io.BytesIO(raw))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    return raw.decode("utf-8", errors="ignore")


@router.post("/upload", response_model=schemas.PolicyDocumentOut)
async def upload_policy(
    title: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: models.User = Depends(get_current_user),
):
    """Upload + immediately ingest (chunk + embed) a policy document."""
    content = await file.read()
    # policy docs are stored under a tenant-level "policies" pseudo-project path
    path = s3_service.save_upload(content, file.filename, f"tenant-{user.tenant_id}-policies")

    policy_doc = models.PolicyDocument(
        tenant_id=user.tenant_id,
        title=title,
        filename=file.filename,
        storage_path=path,
    )
    db.add(policy_doc)
    db.commit()
    db.refresh(policy_doc)

    raw_text = _extract_text(path, file.filename)
    ingest_policy_document(db, policy_doc, raw_text)
    db.refresh(policy_doc)
    return policy_doc


@router.get("", response_model=list[schemas.PolicyDocumentOut])
def list_policies(db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    return db.query(models.PolicyDocument).filter(
        models.PolicyDocument.tenant_id == user.tenant_id
    ).all()


@router.delete("/{policy_id}")
def delete_policy(policy_id: str, db: Session = Depends(get_db), user: models.User = Depends(get_current_user)):
    policy = db.query(models.PolicyDocument).filter(
        models.PolicyDocument.id == policy_id, models.PolicyDocument.tenant_id == user.tenant_id
    ).first()
    if not policy:
        raise HTTPException(404, "Policy document not found")
    db.query(models.PolicyChunk).filter(models.PolicyChunk.document_id == policy.id).delete()
    db.delete(policy)
    db.commit()
    return {"ok": True}


@router.post("/ask", response_model=schemas.PolicyAskResponse)
def ask_policy(payload: schemas.PolicyAskRequest, db: Session = Depends(get_db),
               user: models.User = Depends(get_current_user)):
    """
    Ask a question grounded in the tenant's uploaded policy documents.
    Confirms the RAG pipeline is actually working off what was uploaded —
    if nothing relevant was retrieved, the response says so explicitly
    rather than letting the model improvise an answer.
    """
    has_policies = db.query(models.PolicyDocument).filter(
        models.PolicyDocument.tenant_id == user.tenant_id,
        models.PolicyDocument.ingested == True,  # noqa: E712
    ).count() > 0
    if not has_policies:
        raise HTTPException(400, "No ingested policy documents yet — upload one first")

    result = answer_policy_question(db, user.tenant_id, payload.question)
    return schemas.PolicyAskResponse(**result)
